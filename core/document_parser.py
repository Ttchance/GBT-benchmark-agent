from __future__ import annotations

# -*- coding: utf-8 -*-
"""
文档解析器基类
负责将原始文档（.docx / .pdf）解析为统一的结构化表示，
供各审查器消费。
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
import logging
import re
from core.parse_pdf import GBTPdfParser 

logger = logging.getLogger(__name__)


@dataclass
class Paragraph:
    """段落单元"""
    index: int              # 全文段落序号（从 0 开始）
    section: str            # 所属章节编号，例如 "5.2.1"
    section_title: str      # 章节标题
    text: str               # 纯文本内容
    style: str              # 样式名称（标题/正文/列项 等）
    font_name: str = ""
    font_size: float = 0.0
    is_bold: bool = False
    is_italic: bool = False


@dataclass
class TableCell:
    row: int
    col: int
    text: str


@dataclass
class Table:
    """表格单元"""
    index: int              # 全文表格序号
    caption: str            # 表格标题
    section: str            # 所属章节
    cells: list[TableCell] = field(default_factory=list)


@dataclass
class Figure:
    """图片单元"""
    index: int              # 全文图片序号
    caption: str            # 图片标题
    section: str            # 所属章节
    image_data: bytes = b"" # 原始图片字节（可选）
    image_path: str = ""    # 图片临时路径（可选）


@dataclass
class ParsedDocument:
    """解析后的完整文档对象"""
    file_path: str
    file_format: str                            # "docx" / "pdf"
    paragraphs: list[Paragraph] = field(default_factory=list)
    tables: list[Table] = field(default_factory=list)
    figures: list[Figure] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)    # 文档元数据


class BaseDocumentParser(ABC):
    """文档解析器抽象基类"""

    @abstractmethod
    def parse(self, file_path: str | Path) -> ParsedDocument:
        """
        解析文档并返回结构化的 ParsedDocument。

        Args:
            file_path: 文档路径

        Returns:
            ParsedDocument
        """

    def _check_file(self, file_path: str | Path) -> Path:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")
        if path.suffix.lower() not in (".docx", ".pdf"):
            raise ValueError(f"不支持的文件格式: {path.suffix}")
        return path


class DOCXDocumentParser(BaseDocumentParser):
    """DOCX 文档解析器。"""

    SECTION_RE = re.compile(r"^(\d+(?:\.\d+)*)\s+(.+)$")

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = self._check_file(file_path)

        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("缺少 `python-docx` 依赖，请先执行 `pip install -r requirements.txt`。") from exc

        doc = Document(str(path))
        paragraphs = []
        tables = []
        figures = []
        current_section = ""
        current_section_title = ""

        for index, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else ""
            section_match = self.SECTION_RE.match(text)
            if section_match:
                current_section = section_match.group(1)
                current_section_title = section_match.group(2).strip()

            font_name = ""
            font_size = 0.0
            is_bold = False
            is_italic = False
            for run in para.runs:
                if run.text.strip():
                    font_name = run.font.name or ""
                    font_size = float(run.font.size.pt) if run.font.size else 0.0
                    is_bold = bool(run.bold)
                    is_italic = bool(run.italic)
                    break

            paragraphs.append(
                Paragraph(
                    index=index,
                    section=current_section,
                    section_title=current_section_title,
                    text=text,
                    style=style,
                    font_name=font_name,
                    font_size=font_size,
                    is_bold=is_bold,
                    is_italic=is_italic,
                )
            )

        for table_index, table in enumerate(doc.tables):
            cells = []
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    cells.append(
                        TableCell(
                            row=row_index,
                            col=col_index,
                            text=cell.text.strip(),
                        )
                    )

            tables.append(
                Table(
                    index=table_index,
                    caption="",
                    section="",
                    cells=cells,
                )
            )

        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
            "category": doc.core_properties.category or "",
        }

        return ParsedDocument(
            file_path=str(path),
            file_format="docx",
            paragraphs=paragraphs,
            tables=tables,
            figures=figures,
            metadata=metadata,
        )


class PDFDocumentParser(BaseDocumentParser):
    """PDF 文档解析器。"""

    SECTION_RE = re.compile(r"^(\d+(?:\.\d+)*)\s+(.+)$")

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = self._check_file(file_path)

        try:
            import fitz
        except ImportError as exc:
            raise ImportError("缺少 `pymupdf` 依赖，请先执行 `pip install -r requirements.txt`。") from exc

        doc = fitz.open(str(path))
        paragraphs = []
        tables = []
        figures = []
        current_section = ""
        current_section_title = ""
        paragraph_index = 0

        for page_number, page in enumerate(doc, start=1):
            blocks = page.get_text("blocks")
            for block in blocks:
                text = block[4].strip()
                if not text:
                    continue

                lines = [line.strip() for line in text.splitlines() if line.strip()]
                merged = "\n".join(lines)
                section_match = self.SECTION_RE.match(lines[0]) if lines else None
                if section_match:
                    current_section = section_match.group(1)
                    current_section_title = section_match.group(2).strip()

                paragraphs.append(
                    Paragraph(
                        index=paragraph_index,
                        section=current_section,
                        section_title=current_section_title,
                        text=merged,
                        style="PDFBlock",
                    )
                )
                paragraph_index += 1

            try:
                found_tables = page.find_tables()
            except Exception as exc:
                logger.debug("PDF 表格提取失败: %s", exc)
                found_tables = None

            if found_tables:
                for page_table in found_tables.tables:
                    cells = []
                    extracted = page_table.extract()
                    for row_index, row in enumerate(extracted):
                        for col_index, value in enumerate(row):
                            cells.append(
                                TableCell(
                                    row=row_index,
                                    col=col_index,
                                    text=(value or "").strip(),
                                )
                            )
                    tables.append(
                        Table(
                            index=len(tables),
                            caption="",
                            section=current_section,
                            cells=cells,
                        )
                    )

        metadata = {
            "title": doc.metadata.get("title", "") if doc.metadata else "",
            "author": doc.metadata.get("author", "") if doc.metadata else "",
            "page_count": doc.page_count,
        }
        doc.close()

        return ParsedDocument(
            file_path=str(path),
            file_format="pdf",
            paragraphs=paragraphs,
            tables=tables,
            figures=figures,
            metadata=metadata,
        )


class AutoDocumentParser(BaseDocumentParser):
    """根据文件后缀自动选择解析器。"""

    # def __init__(self, llm_client=None):
    def __init__(self, llm_client):
        self._llm_client = llm_client

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = self._check_file(file_path)
        if path.suffix.lower() == ".docx":
            return DOCXDocumentParser().parse(path)
        return GBTPdfParser(llm_client=self._llm_client).parse(path)
